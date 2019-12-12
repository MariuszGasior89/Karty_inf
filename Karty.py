#!/usr/bin/env python
# -*- coding: utf-8 -*-

# skrypt do tworzenia kart informacyjnych.
import os

import arcpy
import docx
from docx.shared import Cm, Pt

from gdb2list import listaDrog, listaElek, listaPob, listaOch, listaLch, listaOc, listaZoo, listaCm, listaKap, listaOcz, listaUj, listaSkl, listaZak, listaMel, listaKru, listaPrz, listaGor

arcpy.env.overwriteOutput = True
arcpy.env.workspace = ws = ur'C:\ROBOCZY\skrypty_python\Karty_inf\Baza_do_ankiet\Baza_do_ankiet.gdb'
aJCWP = 'JCWP_rzeczne'
stareJCWP = 'Status_JCWP_rzeczne'

fold = ur"C:\ROBOCZY\skrypty_python\Karty_inf\Mapy_karty"
kol = []

for field in arcpy.ListFields(aJCWP):
        kol.append(str(field.name))

lista_drog = []
listaDrog(lista_drog)

lista_elek = []
listaElek(lista_elek)

lista_pob = []
listaPob(lista_pob)

lista_och = []
listaOch(lista_och)

lista_lch = []
listaLch(lista_lch)

lista_oc = []
listaOc(lista_oc)

lista_zoo = []
listaZoo(lista_zoo)

lista_cm = []
listaCm(lista_cm)

lista_kap = []
listaKap(lista_kap)

lista_oczysz = []
listaOcz(lista_oczysz)

lista_uj = []
listaUj(lista_uj)

lista_sklad = []
listaSkl(lista_sklad)

lista_zaklad = []
listaZak(lista_zaklad)

lista_mel = []
listaMel(lista_mel)

lista_krusz = []
listaKru(lista_krusz)

lista_prze = []
listaPrz(lista_prze)

lista_gor = []
listaGor(lista_gor)


def mapa(fold_map, jcwp):

    for plik in os.listdir(fold_map):
        if os.path.isfile(os.path.join(fold_map, plik)) and os.path.splitext(plik)[1] == ".jpg" \
                and os.path.splitext(plik)[0] == jcwp:
            mapka = os.path.join(fold_map, plik)
    return mapka


def add_naglowek(tyt_nag):

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    paragraph = paragraph.add_run(tyt_nag)
    font = paragraph.font
    font.name = 'Arial'
    font.bold = True
    font.size = Pt(10)
    return


def add_objasnienia(tyt_nag, bold=True):

    paragraph = document.add_paragraph()
    paragraph = paragraph.add_run(tyt_nag)
    font = paragraph.font
    font.name = 'Arial'
    font.bold = bold
    font.size = Pt(10)
    return


def add_tabhead(tab_head, paragraph):

    run = paragraph.add_run(tab_head)
    font = run.font
    font.name = 'Arial'
    font.bold = True
    font.size = Pt(10)
    return


def add_run_bold(tab_run, paragraph, bold=True):

    run = paragraph.add_run(tab_run)
    font = run.font
    font.name = 'Arial'
    font.bold = bold
    font.size = Pt(10)
    return


def add_tabrun(tab_run, paragraph):

    run = paragraph.add_run(tab_run)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(10)
    return


with arcpy.da.UpdateCursor(aJCWP, kol) as cur:
    for row in cur:
        # warunek filtrujacy aby generowac warstwy tylko dla wybranych zlewni
        if row[kol.index('KARTA')] == 999:
            print(row[kol.index('MS_KOD')])
            document = docx.Document()

            """Dodanie nagłowka 1"""
            add_naglowek('Karta informacyjna aJCWP')

            """Dodanie Tabeli 1"""
            tab1 = document.add_table(rows=4, cols=2, style='Table Grid')

            for cell in tab1.columns[0].cells:
                cell.width = Cm(3)

            for cell in tab1.columns[1].cells:
                cell.width = Cm(12)

            add_tabhead('Kod', tab1.cell(0, 0).paragraphs[0])
            add_tabhead('Nazwa', tab1.cell(1, 0).paragraphs[0])
            add_tabhead('Typ abiotyczny', tab1.cell(2, 0).paragraphs[0])
            add_tabhead(u'Długość [km]', tab1.cell(3, 0).paragraphs[0])

            add_tabrun(row[kol.index('MS_KOD')], tab1.cell(0, 1).paragraphs[0])
            add_tabrun(row[kol.index('Nazwa_JCWP')], tab1.cell(1, 1).paragraphs[0])
            add_tabrun(row[kol.index('TYP_ABIOT')], tab1.cell(2, 1).paragraphs[0])
            add_tabrun(str(row[kol.index('DLUG_RZ')]), tab1.cell(3, 1).paragraphs[0])
            document.add_paragraph()

            """Dodanie nagłowka 2"""
            add_naglowek('Informacja o JCWP z aPGW')

            """Dodanie Tabeli 2"""
            wiersze = row[kol.index('IL_JCWP_aPGW')]
            tab2 = document.add_table(rows=5+(2*wiersze), cols=4, style='Table Grid')

            for cell in tab2.columns[0].cells:
                cell.width = Cm(4)

            for cell in tab2.columns[1].cells:
                cell.width = Cm(1.5)

            for cell in tab2.columns[2].cells:
                cell.width = Cm(1.5)

            for cell in tab2.columns[3].cells:
                cell.width = Cm(8)

            add_tabhead('Relacja aJCWP do JCWP', tab2.cell(0, 0).paragraphs[0])
            add_tabrun('1:' + str(wiersze), tab2.cell(0, 2).paragraphs[0])
            add_tabhead('Wynik wyznaczenia w II cyklu planistycznym', tab2.cell(1, 2).paragraphs[0])
            add_tabhead('Kod JCWP', tab2.cell(2, 0).paragraphs[0])
            add_tabhead(u'Udział % w aJCWP', tab2.cell(2, 1).paragraphs[0])
            add_tabhead('Status', tab2.cell(2, 2).paragraphs[0])
            add_tabhead(u'Zmiany hydromorfologiczne uzasadniające wyznaczenie', tab2.cell(2, 3).paragraphs[0])

            add_tabhead(u'Ocena stanu i perspektyw osiągniecia  celów dla stanu/potencjału ekologicznego w II cyklu planistycznym ', tab2.cell(wiersze + 3, 1).paragraphs[0])
            add_tabhead('Kod JCWP', tab2.cell(wiersze + 4, 0).paragraphs[0])
            add_tabhead(u'Stan/potencjał', tab2.cell(wiersze + 4, 1).paragraphs[0])
            add_tabhead('Ocena ryzyka', tab2.cell(wiersze + 4, 2).paragraphs[0])
            add_tabhead(u'Cel dla stanu/potencjału ekologicznego', tab2.cell(wiersze + 4, 3).paragraphs[0])

            old_JCWP_list = [row[kol.index('ST_KOD_1')], row[kol.index('ST_KOD_2')], row[kol.index('ST_KOD_3')],row[kol.index('ST_KOD_4')],row[kol.index('ST_KOD_5')],row[kol.index('ST_KOD_6')],row[kol.index('ST_KOD_7')],row[kol.index('ST_KOD_8')],row[kol.index('ST_KOD_9')], row[kol.index('ST_KOD_10')],row[kol.index('ST_KOD_11')],row[kol.index('ST_KOD_12')],row[kol.index('ST_KOD_13')],row[kol.index('ST_KOD_14')],row[kol.index('ST_KOD_15')]]

            j=0
            kol_stare = []
            for field in arcpy.ListFields(stareJCWP):
                kol_stare.append(str(field.name))

            for old_JCWP in old_JCWP_list[0:wiersze]:
                 with arcpy.da.SearchCursor(stareJCWP, kol_stare) as cur_stare:
                    for row_stare in cur_stare:
                        if old_JCWP == row_stare[kol_stare.index('KOD_JCWP')]:
                            add_tabrun(str(row_stare[kol_stare.index('KOD_JCWP')]), tab2.cell(3 + j, 0).paragraphs[0])
                            add_tabrun(str(row_stare[kol_stare.index('UDZIAL')]), tab2.cell(3 + j, 1).paragraphs[0])
                            add_tabrun(str(row_stare[kol_stare.index('Status_os')]), tab2.cell(3 + j, 2).paragraphs[0])
                            add_tabrun(unicode(row_stare[kol_stare.index('UZAS')]), tab2.cell(3 + j, 3).paragraphs[0])

                            add_tabrun(str(row_stare[kol_stare.index('KOD_JCWP')]), tab2.cell(5 + j + wiersze, 0).paragraphs[0])
                            add_tabrun(unicode(row_stare[kol_stare.index('Stan_pot')]), tab2.cell(5 + j + wiersze, 1).paragraphs[0])
                            add_tabrun(unicode(row_stare[kol_stare.index('Ryzyko')]), tab2.cell(5 + j + wiersze, 2).paragraphs[0])
                            add_tabrun(unicode(row_stare[kol_stare.index('Cel_Stan_pot')]), tab2.cell(5 + j + wiersze, 3).paragraphs[0])
                    j+=1
                 del cur_stare

            tab2.cell(0, 0).merge(tab2.cell(0, 1))
            tab2.cell(0, 2).merge(tab2.cell(0, 3))
            tab2.cell(1, 0).merge(tab2.cell(1, 1))
            tab2.cell(1, 2).merge(tab2.cell(1, 3))
            tab2.cell(wiersze+3, 1).merge(tab2.cell(wiersze+3, 3))
            document.add_paragraph()

            """Dodanie nagłowka 3"""
            add_naglowek('Mapa')

            """Dodanie Mapy"""
            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            paragraph = paragraph.add_run()
            paragraph.add_picture(mapa(fold, row[kol.index('MS_KOD')]), width=Cm(12.5))
            document.add_paragraph()

            """Dodanie nagłowka 4"""
            add_naglowek(u'Wyniki wstępnego wyznaczania w III cyklu planistycznym')

            """Dodanie Tabeli 3"""
            tab3 = document.add_table(rows=12, cols=6, style='Table Grid')
            add_tabhead(u'Ciek główny', tab3.cell(0, 1).paragraphs[0])
            add_tabhead(u'Cieki pozostałe', tab3.cell(0, 2).paragraphs[0])
            add_tabhead(u'Ciek główny', tab3.cell(0, 4).paragraphs[0])
            add_tabhead(u'Cieki pozostałe', tab3.cell(0, 5).paragraphs[0])
            add_tabhead('PRH2 (0-5)', tab3.cell(1, 0).paragraphs[0])
            add_tabhead('PRH3 (0-5)', tab3.cell(2, 0).paragraphs[0])
            add_tabhead('PRH4 (1-10)', tab3.cell(3, 0).paragraphs[0])
            add_tabhead('PRH5 (0-5)', tab3.cell(4, 0).paragraphs[0])
            add_tabhead('PRH6 (0-5)', tab3.cell(5, 0).paragraphs[0])
            add_tabhead('PRH7 (0-5)', tab3.cell(6, 0).paragraphs[0])
            add_tabhead('WRH', tab3.cell(9, 0).paragraphs[0])
            add_tabhead('PPH2 (0-8)', tab3.cell(1, 3).paragraphs[0])
            add_tabhead('PPH3 (0-6)', tab3.cell(2, 3).paragraphs[0])
            add_tabhead('PPH4 (0-5)', tab3.cell(3, 3).paragraphs[0])
            add_tabhead('PPH5 (0-3)', tab3.cell(4, 3).paragraphs[0])
            add_tabhead('PPH6 (0-7)', tab3.cell(5, 3).paragraphs[0])
            add_tabhead('PPH7 (0-6)', tab3.cell(6, 3).paragraphs[0])
            add_tabhead('WPH', tab3.cell(9, 3).paragraphs[0])
            add_tabhead(u'Parametry różnorodności hydromorfologicznej', tab3.cell(0, 0).paragraphs[0])
            add_tabhead(u'Parametry przekształcenia hydromorfologicznego', tab3.cell(0, 3).paragraphs[0])
            add_tabhead('WTR', tab3.cell(8, 0).paragraphs[0])
            add_tabhead('WPTR', tab3.cell(8, 3).paragraphs[0])
            add_tabhead('HIRk', tab3.cell(10, 0).paragraphs[0])
            add_tabhead(u'Status wstępny', tab3.cell(11, 0).paragraphs[0])

            tab3.cell(7, 0).merge(tab3.cell(7, 5))
            tab3.cell(10, 3).merge(tab3.cell(10, 5))
            tab3.cell(11, 3).merge(tab3.cell(11, 5))

            add_tabrun(str(row[kol.index('PRH2_G')]), tab3.cell(1, 1).paragraphs[0])
            add_tabrun(str(row[kol.index('PRH3_G')]), tab3.cell(2, 1).paragraphs[0])
            add_tabrun(str(row[kol.index('PRH4_G')]), tab3.cell(3, 1).paragraphs[0])
            add_tabrun(str(row[kol.index('PRH5_G')]), tab3.cell(4, 1).paragraphs[0])
            add_tabrun(str(row[kol.index('PRH6_G')]), tab3.cell(5, 1).paragraphs[0])
            add_tabrun(str(row[kol.index('PRH7_G')]), tab3.cell(6, 1).paragraphs[0])
            add_tabrun(str(row[kol.index('PRH2_P')]), tab3.cell(1, 2).paragraphs[0])
            add_tabrun(str(row[kol.index('PRH3_P')]), tab3.cell(2, 2).paragraphs[0])
            add_tabrun(str(row[kol.index('PRH4_P')]), tab3.cell(3, 2).paragraphs[0])
            add_tabrun(str(row[kol.index('PRH5_P')]), tab3.cell(4, 2).paragraphs[0])
            add_tabrun(str(row[kol.index('PRH6_P')]), tab3.cell(5, 2).paragraphs[0])
            add_tabrun(str(row[kol.index('PRH7_P')]), tab3.cell(6, 2).paragraphs[0])
            add_tabrun(str(row[kol.index('PPH2_G')]), tab3.cell(1, 4).paragraphs[0])
            add_tabrun(str(row[kol.index('PPH3_G')]), tab3.cell(2, 4).paragraphs[0])
            add_tabrun(str(row[kol.index('PPH4_G')]), tab3.cell(3, 4).paragraphs[0])
            add_tabrun(str(row[kol.index('PPH5_G')]), tab3.cell(4, 4).paragraphs[0])
            add_tabrun(str(row[kol.index('PPH6_G')]), tab3.cell(5, 4).paragraphs[0])
            add_tabrun(str(row[kol.index('PPH7_G')]), tab3.cell(6, 4).paragraphs[0])
            add_tabrun(str(row[kol.index('PPH2_P')]), tab3.cell(1, 5).paragraphs[0])
            add_tabrun(str(row[kol.index('PPH3_P')]), tab3.cell(2, 5).paragraphs[0])
            add_tabrun(str(row[kol.index('PPH4_P')]), tab3.cell(3, 5).paragraphs[0])
            add_tabrun(str(row[kol.index('PPH5_P')]), tab3.cell(4, 5).paragraphs[0])
            add_tabrun(str(row[kol.index('PPH6_P')]), tab3.cell(5, 5).paragraphs[0])
            add_tabrun(str(row[kol.index('PPH7_P')]), tab3.cell(6, 5).paragraphs[0])

            add_tabrun(str(row[kol.index('WRH_G')]), tab3.cell(9, 1).paragraphs[0])
            add_tabrun(str(row[kol.index('WTR_G')]), tab3.cell(8, 1).paragraphs[0])
            add_tabrun(str(row[kol.index('WPTR_G')]), tab3.cell(8, 4).paragraphs[0])
            add_tabrun(str(row[kol.index('WRH_P')]), tab3.cell(9, 2).paragraphs[0])
            add_tabrun(str(row[kol.index('WTR_P')]), tab3.cell(8, 2).paragraphs[0])
            add_tabrun(str(row[kol.index('WPTR_P')]), tab3.cell(8, 5).paragraphs[0])
            add_tabrun(str(row[kol.index('WPH_G')]), tab3.cell(9, 4).paragraphs[0])
            add_tabrun(str(row[kol.index('WPH_P')]), tab3.cell(9, 5).paragraphs[0])
            add_tabrun(str(row[kol.index('HIRk')]), tab3.cell(10, 1).paragraphs[0])
            add_tabhead(str(row[kol.index('WYZN_WST')]), tab3.cell(11, 1).paragraphs[0])
            tab3.cell(10, 1).merge(tab3.cell(10, 2))
            tab3.cell(11, 1).merge(tab3.cell(11, 2))

            document.add_paragraph()
            """Dodanie OBJAŚNIENIA"""
            add_objasnienia(u'Objaśnienia:', False)
            add_objasnienia(u'Parametry Różnorodności Hydromorfologicznej (PRH):', True)
            par_obj1 = document.add_paragraph()
            add_run_bold(u'PRH2 - ', par_obj1, True)
            add_run_bold(u'Odsypy śródkorytowe i wyspy, ', par_obj1, False)
            add_run_bold(u'PRH3 - ', par_obj1, True)
            add_run_bold(u'Odsypy brzegowe, ', par_obj1, False)
            add_run_bold(u'PRH4 - ', par_obj1, True)
            add_run_bold(u'Użytkowanie terenu doliny rzecznej, ', par_obj1, False)
            add_run_bold(u'PRH5 - ', par_obj1, True)
            add_run_bold(u'Zadrzewienia, ', par_obj1, False)
            add_run_bold(u'PRH6 - ', par_obj1, True)
            add_run_bold(u'Starorzecza i inne niewielkie zbiorniki wodne, ', par_obj1, False)
            add_run_bold(u'PRH7 - ', par_obj1, True)
            add_run_bold(u'Tereny podmokłe', par_obj1, False)

            add_objasnienia(u'Parametry Przekształcenia Hydromorfologicznego (PPH):', True)
            par_obj2 = document.add_paragraph()
            add_run_bold(u'PPH2 - ', par_obj2, True)
            add_run_bold(u'Budowle piętrzące, ', par_obj2, False)
            add_run_bold(u'PPH3 - ', par_obj2, True)
            add_run_bold(u'Obiekty gospodarki wodnej, ', par_obj2, False)
            add_run_bold(u'PPH4 - ', par_obj2, True)
            add_run_bold(u'Budowle regulacyjne, ', par_obj2, False)
            add_run_bold(u'PPH5 - ', par_obj2, True)
            add_run_bold(u'Obiekty mostowe i przeprawy, ', par_obj2, False)
            add_run_bold(u'PPH6 - ', par_obj2, True)
            add_run_bold(u'Obwałowania, ', par_obj2, False)
            add_run_bold(u'PPH7 - ', par_obj2, True)
            add_run_bold(u'Wpływ górnictwa', par_obj2, False)

            par_obj3 = document.add_paragraph()
            add_run_bold(u'WTR - ', par_obj3, True)
            add_run_bold(u'Wskaźnik trasy rzeki, ', par_obj3, False)
            add_run_bold(u'WRH - ', par_obj3, True)
            add_run_bold(u'Wskaźnik różnorodności hydromorfologicznej, ', par_obj3, False)
            add_run_bold(u'WPTR - ', par_obj3, True)
            add_run_bold(u'Wskaźnik przekształcenia trasy rzeki, ', par_obj3, False)
            add_run_bold(u'WPH - ', par_obj3, True)
            add_run_bold(u'Wskaźnik przekształcenia hydromorfologii, ', par_obj3, False)
            add_run_bold(u'HIR - ', par_obj3, True)
            add_run_bold(u'Hydromorfologiczny Indeks Rzeczny', par_obj3, False)
            document.add_paragraph()

            """Dodanie nagłowka 5"""
            add_naglowek(u'Korekta ekspercka')

            """Dodanie Tabeli 4"""
            tab4 = document.add_table(rows=3, cols=2, style='Table Grid')

            for cell in tab4.columns[0].cells:
                cell.width = Cm(5.5)

            for cell in tab4.columns[1].cells:
                cell.width = Cm(9.5)

            add_tabhead(u'Czy dokonano korekty?', tab4.cell(0, 0).paragraphs[0])
            add_tabhead(u'Opis procedury/uzasadnienie', tab4.cell(1, 0).paragraphs[0])
            add_tabhead(u'Status wstępny skorygowany', tab4.cell(2, 0).paragraphs[0])

            add_tabrun(str(row[kol.index('KOREKT_EKS')]), tab4.cell(0, 1).paragraphs[0])
            add_tabrun(unicode(row[kol.index('UZAS')]), tab4.cell(1, 1).paragraphs[0])

            if row[kol.index('WYZN_WST')] == u'NAT':
                add_tabrun("SZCW", tab4.cell(2, 1).paragraphs[0])

            if row[kol.index('WYZN_WST')] == u'SZCW':
                add_tabrun("SCW", tab4.cell(2, 1).paragraphs[0])

            document.add_paragraph()

            """Dodanie nagłowka 6"""
            add_naglowek(u'Znaczące zmiany hydromorfologiczne')

# wyznaczenie wspolczynnikow do okreslenia znaczacych zmian hydromorfologicznych
            par_ind = []
            PPH2 = int(row[kol.index('PPH2_G')] + row[kol.index('PPH2_P')])
            PPH3 = int(row[kol.index('PPH3_G')] + row[kol.index('PPH3_P')])
            PPH4 = int(row[kol.index('PPH4_G')] + row[kol.index('PPH4_P')])
            PPH5 = int(row[kol.index('PPH5_G')] + row[kol.index('PPH5_P')])
            PPH6 = row[kol.index('PPH6_G')] + row[kol.index('PPH6_P')]
            PPH7 = int(row[kol.index('PPH7_G')] + row[kol.index('PPH7_P')])

            if PPH2 > 0:
                par_ind.append('PPH2')
            if PPH3 > 0:
                par_ind.append('PPH3')
            if PPH4 > 0:
                par_ind.append('PPH4')
            if PPH5 > 0:
                par_ind.append('PPH5')
            if PPH6 > 0:
                par_ind.append('PPH6')
            if PPH7 > 0:
                par_ind.append('PPH7')

            """Dodanie Tabeli 5"""
            tab5 = document.add_table(rows = (len(par_ind)+1), cols=2, style='Table Grid')
            for cell in tab5.columns[0].cells:
                cell.width = Cm(2.5)

            for cell in tab5.columns[1].cells:
                cell.width = Cm(12.5)
            i=0
            for par in par_ind:
                add_tabhead(par, tab5.cell(i, 0).paragraphs[0])
                kol2=[]
                for field in arcpy.ListFields(par):
                    kol2.append(str(field.name))
                with arcpy.da.SearchCursor(par, kol2) as cur_pph2:
                    for row_pph2 in cur_pph2:
                        if row[kol.index('MS_KOD')] == row_pph2[kol2.index('MS_KOD')] and PPH2 > 0 and par == 'PPH2':
                            if row[kol.index('PPH2_G')] > 0  and row[kol.index('PPH2_P')] > 0:
                                add_tabrun(u"Istotne budowle piętrzące w JCWP - {0} - {1} budowli na km cieku głównego oraz {2} budowli na km cieków pozostałych.".format(int(row_pph2[kol2.index('OB_ODDZ')]), row_pph2[kol2.index('RZEKI_GL')], row_pph2[kol2.index('RZEKI_POZ')]), tab5.cell(i, 1).paragraphs[0])
                            elif row[kol.index('PPH2_G')] > 0  and row[kol.index('PPH2_P')] == 0:
                                add_tabrun(u"Istotne budowle piętrzące w JCWP - {0} - {1} budowli na km cieku głównego.".format(int(row_pph2[kol2.index('OB_ODDZ')]), row_pph2[kol2.index('RZEKI_GL')]), tab5.cell(i, 1).paragraphs[0])
                            elif row[kol.index('PPH2_G')] == 0  and row[kol.index('PPH2_P')] > 0:
                                add_tabrun(u"Istotne budowle piętrzące w JCWP - {0} - {1} budowli na km cieków pozostałych.".format(int(row_pph2[kol2.index('OB_ODDZ')]), row_pph2[kol2.index('RZEKI_POZ')]), tab5.cell(i, 1).paragraphs[0])

                        elif row[kol.index('MS_KOD')] == row_pph2[kol2.index('MS_KOD')] and PPH3 > 0 and par == 'PPH3':
                            if row[kol.index('PPH3_G')] > 0  and row[kol.index('PPH3_P')] > 0:
                                add_tabrun(u"Udział obiektów gospodarki wodnej w powierzchni najbliższego sąsiedztwa cieku głównego wynosi {0}% oraz {1}% dla powierzchni najbliższego sąsiedztwa cieków pozostałych w JCWP.".format(row_pph2[kol2.index('UDZ_SZT_ZB')], row_pph2[kol2.index('UDZ_SZT_ZB_1')]), tab5.cell(i, 1).paragraphs[0])
                            elif row[kol.index('PPH3_G')] > 0  and row[kol.index('PPH3_P')] == 0:
                                add_tabrun(u"Udział obiektów gospodarki wodnej w powierzchni najbliższego sąsiedztwa cieku głównego wynosi {0}% w JCWP.".format(row_pph2[kol2.index('UDZ_SZT_ZB')]), tab5.cell(i, 1).paragraphs[0])
                            elif row[kol.index('PPH3_G')] == 0  and row[kol.index('PPH3_P')] > 0:
                                add_tabrun(u"Udział obiektów gospodarki wodnej w powierzchni najbliższego sąsiedztwa cieków pozostałych wynosi {0}% w JCWP.".format(row_pph2[kol2.index('UDZ_SZT_ZB_1')]), tab5.cell(i, 1).paragraphs[0])

                        elif row[kol.index('MS_KOD')] == row_pph2[kol2.index('MS_KOD')] and PPH4 > 0 and par == 'PPH4':
                            if row[kol.index('PPH4_G')] > 0  and row[kol.index('PPH4_P')] > 0:
                                add_tabrun(u"Udział ciężkich budowli regulacyjnych -{0}% i lekkich budowli regulacyjnych -{1}% w stosunku do długości cieku głównego JCWP. Udział ciężkich budowli regulacyjnych -{2}% i lekkich budowli regulacyjnych -{3}% w stosunku do długości cieków pozostałych JCWP.".format(row_pph2[kol2.index('UM_C')], row_pph2[kol2.index('UM_L')], row_pph2[kol2.index('UM_C_1')], row_pph2[kol2.index('UM_L_1')]), tab5.cell(i, 1).paragraphs[0])
                            elif row[kol.index('PPH4_G')] > 0  and row[kol.index('PPH4_P')] == 0:
                                add_tabrun(u"Udział ciężkich budowli regulacyjnych -{0}% i lekkich budowli regulacyjnych -{1}% w stosunku do długości cieku głównego JCWP.".format(row_pph2[kol2.index('UM_C')], row_pph2[kol2.index('UM_L')]), tab5.cell(i, 1).paragraphs[0])
                            elif row[kol.index('PPH4_G')] == 0  and row[kol.index('PPH4_P')] > 0:
                                add_tabrun(u"Udział ciężkich budowli regulacyjnych -{0}% i lekkich budowli regulacyjnych -{1}% w stosunku do długości cieków pozostałych JCWP.".format(row_pph2[kol2.index('UM_C_1')], row_pph2[kol2.index('UM_L_1')]), tab5.cell(i, 1).paragraphs[0])


                        elif row[kol.index('MS_KOD')] == row_pph2[kol2.index('MS_KOD')] and PPH5 > 0 and par == 'PPH5':
                            if row[kol.index('PPH5_G')] > 0  and row[kol.index('PPH5_P')] > 0:
                                add_tabrun(u"Obiekty mostowe i przeprawy w JCWP – {0} – {1} obiektów na km cieku głównego oraz {2} obiektów na km cieków pozostałych.".format(int(row_pph2[kol2.index('IL_PRZEPRAW')]), row_pph2[kol2.index('L_MOST_KM')], row_pph2[kol2.index('L_MOST_KM_1')]), tab5.cell(i, 1).paragraphs[0])
                            elif row[kol.index('PPH5_G')] > 0  and row[kol.index('PPH5_P')] == 0:
                                add_tabrun(u"Obiekty mostowe i przeprawy w JCWP – {0} – {1} obiektów na km cieku głównego.".format(int(row_pph2[kol2.index('IL_PRZEPRAW')]), row_pph2[kol2.index('L_MOST_KM')]), tab5.cell(i, 1).paragraphs[0])
                            elif row[kol.index('PPH5_G')] == 0  and row[kol.index('PPH5_P')] > 0:
                                add_tabrun(u"Obiekty mostowe i przeprawy w JCWP – {0} – {1} obiektów na km cieków pozostałych.".format(int(row_pph2[kol2.index('IL_PRZEPRAW')]), row_pph2[kol2.index('L_MOST_KM_1')]), tab5.cell(i, 1).paragraphs[0])

                        elif row[kol.index('MS_KOD')] == row_pph2[kol2.index('MS_KOD')] and PPH6 > 0 and par == 'PPH6':
                            if (row[kol.index('PPH6_G')] > 0  and row[kol.index('PPH6_P')] > 0):
                                add_tabrun(u"Długość obwałowanych cieków w JCWP wynosi {0}km. Ciek główny obwałowano w {1}% dwustronnie i {2}% jednostronnie. Międzywale na cieku głównym w przeważającej części – {3}. Cieki pozostałe obwałowano w {4}% dwustronnie i {5}% jednostronnie. Międzywale na ciekach pozostałych w przeważającej części – {6}.".format(row_pph2[kol2.index('DL_OBW')], row_pph2[kol2.index('UDZ_OBW_DWU')], row_pph2[kol2.index('UDZ_OBW_JED')], row_pph2[kol2.index('MIEDZYWAL')], row_pph2[kol2.index('UDZ_OBW_DWU_1')], row_pph2[kol2.index('UDZ_OBW_JED_1')], row_pph2[kol2.index('MIEDZYWAL_1')]), tab5.cell(i, 1).paragraphs[0])
                            elif (row[kol.index('PPH6_G')] > 0  and row[kol.index('PPH6_P')] == 0):
                                add_tabrun(u"Długość obwałowanych cieków w JCWP wynosi {0}km. Ciek główny obwałowano w {1}% dwustronnie i {2}% jednostronnie. Międzywale na cieku głównym w przeważającej części – {3}.".format(row_pph2[kol2.index('DL_OBW')], row_pph2[kol2.index('UDZ_OBW_DWU')], row_pph2[kol2.index('UDZ_OBW_JED')], row_pph2[kol2.index('MIEDZYWAL')]), tab5.cell(i, 1).paragraphs[0])
                            elif (row[kol.index('PPH6_G')] == 0  and row[kol.index('PPH6_P')] > 0):
                                add_tabrun(u"Długość obwałowanych cieków w JCWP wynosi {0}km. Cieki pozostałe obwałowano w {1}% dwustronnie i {2}% jednostronnie. Międzywale na ciekach pozostałych w przeważającej części – {3}.".format(row_pph2[kol2.index('DL_OBW')], row_pph2[kol2.index('UDZ_OBW_DWU_1')], row_pph2[kol2.index('UDZ_OBW_JED_1')], row_pph2[kol2.index('MIEDZYWAL_1')]), tab5.cell(i, 1).paragraphs[0])

                        elif row[kol.index('MS_KOD')] == row_pph2[kol2.index('MS_KOD')] and PPH7 > 0 and par == 'PPH7':
                            if row[kol.index('PPH7_G')] > 0  and row[kol.index('PPH7_P')] > 0:
                                add_tabrun(u"Udział obszarów objętych wpływem działalności górniczej w powierzchni najbliższego sąsiedztwa cieku głównego wynosi {0}% oraz {1}% dla powierzchni najbliższego sąsiedztwa cieków pozostałych w JCWP.".format(row_pph2[kol2.index('UDZ_GORN')], row_pph2[kol2.index('UDZ_GORN_1')]), tab5.cell(i, 1).paragraphs[0])
                            elif row[kol.index('PPH7_G')] > 0  and row[kol.index('PPH7_P')] == 0:
                                add_tabrun(u"Udział obszarów objętych wpływem działalności górniczej w powierzchni najbliższego sąsiedztwa cieku głównego wynosi {0}% w JCWP.".format(row_pph2[kol2.index('UDZ_GORN')]), tab5.cell(i, 1).paragraphs[0])
                            elif row[kol.index('PPH7_G')] == 0  and row[kol.index('PPH7_P')] > 0:
                                add_tabrun(u"Udział obszarów objętych wpływem działalności górniczej w powierzchni najbliższego sąsiedztwa cieków pozostałych wynosi {0}% w JCWP.".format(row_pph2[kol2.index('UDZ_GORN_1')]), tab5.cell(i, 1).paragraphs[0])

                i+=1

                del cur_pph2

            add_tabhead(u'Komentarz', tab5.cell(len(par_ind), 0).paragraphs[0])

            document.add_paragraph()

            """Dodanie nagłowka 7"""
            add_naglowek(u'Informacje o aJCWP z dostępnych baz danych')

            """Dodanie Tabeli 6"""
            tab6 = document.add_table(rows=12, cols=2, style='Table Grid')

            for cell in tab6.columns[0].cells:
                cell.width = Cm(10)

            for cell in tab6.columns[1].cells:
                cell.width = Cm(5)

            add_tabhead(u'Aktualna klasa drogi wodnej', tab6.cell(0, 0).paragraphs[0])
            add_tabhead(u'Wykorzystanie energii wody – elektrownie wodne', tab6.cell(1, 0).paragraphs[0])
            add_tabhead(u'Wielkość poborów  z wód powierzchniowych, dopuszczona w PWP, wg celu poboru [m3/rok]:', tab6.cell(2, 0).paragraphs[0])
            for i in range(9):
                tab6.cell(2, 0).add_paragraph()

            add_tabhead(u'1) rolnictwo-nawodnienia', tab6.cell(2, 0).paragraphs[2])
            add_tabhead(u'2) zaopatrzenie ludności w wodę do spożycia', tab6.cell(2, 0).paragraphs[3])
            add_tabhead(u'3) przemysł', tab6.cell(2, 0).paragraphs[4])
            add_tabhead(u'4) elektrownie wodne', tab6.cell(2, 0).paragraphs[5])
            add_tabhead(u'5) wody chłodnicze', tab6.cell(2, 0).paragraphs[6])
            add_tabhead(u'6) akwakultura', tab6.cell(2, 0).paragraphs[7])
            add_tabhead(u'7) inne', tab6.cell(2, 0).paragraphs[8])

            add_tabhead(u'Obszar, na którym ryzyko powodziowe jest minimalizowane - ISOK [ha]', tab6.cell(3, 0).paragraphs[0])
            add_tabhead(u'Liczba ludzi zamieszkałych w chronionym obszarze - ISOK', tab6.cell(4, 0).paragraphs[0])
            add_tabhead(u'Ważniejsze obiekty kulturowe, zagrażające środowisku w razie powodzi, które są chronione', tab6.cell(5, 0).paragraphs[0])
            add_tabhead(u'Ważniejsze obiekty infrastruktury społecznej, zagrażające środowisku w razie powodzi, które są chronione', tab6.cell(6, 0).paragraphs[0])
            add_tabhead(u'Ważniejsze obiekty gospodarcze, zagrażające środowisku w razie powodzi, które są chronione', tab6.cell(7, 0).paragraphs[0])
            add_tabhead(u'Obszary zmeliorowane [ha]', tab6.cell(8, 0).paragraphs[0])
            add_tabhead(u'Pobory kruszywa [m]', tab6.cell(9, 0).paragraphs[0])
            add_tabhead(u'Przerzuty wody', tab6.cell(10, 0).paragraphs[0])
            add_tabhead(u'Wpływ górnictwa:', tab6.cell(11, 0).paragraphs[0])
            for i in range(3):
                tab6.cell(11, 0).add_paragraph()
            for i in range(4):
                tab6.cell(11, 1).add_paragraph()
            add_tabhead(u'1) powierzchnia terenów górniczych w zlewni jcwp [ha]', tab6.cell(11, 0).paragraphs[1])
            add_tabhead(u'2) powierzchnia obszarów górniczych w zlewni jcwp [ha]', tab6.cell(11, 0).paragraphs[2])
            add_tabhead(u'3) odwodnienia kopalni, sumaryczny pobór [m3/rok]', tab6.cell(11, 0).paragraphs[3])

            #Aktualna klasa drogi wodnej
            if row[kol.index('MS_KOD')] not in lista_drog:
                add_tabrun(u'brak drogi wodnej', tab6.cell(0, 1).paragraphs[0])
            else:
                a1 = a2 = a3 = a4 = a5= a6 = a7 = ''
                with arcpy.da.SearchCursor('drogi_wodne', ['MS_KOD','klasa_drog','SUM_dlg_km']) as cur_drogi:
                    for row_drogi in cur_drogi:
                        if  row[kol.index('MS_KOD')] == row_drogi[0]:
                            if row_drogi[1] == 'Ia':
                                a1 = row_drogi[1] +' - '+str(row_drogi[2])+'km'
                            if row_drogi[1] == 'Ib':
                                a2 = row_drogi[1] +' - '+str(row_drogi[2])+'km'
                            if row_drogi[1] == 'II':
                                a3 = row_drogi[1] +' - '+str(row_drogi[2])+'km'
                            if row_drogi[1] == 'III':
                                a4 = row_drogi[1] +' - '+str(row_drogi[2])+'km'
                            if row_drogi[1] == 'IV':
                                a5 = row_drogi[1] +' - '+str(row_drogi[2])+'km'
                            if row_drogi[1] == 'Va':
                                a6 = row_drogi[1] +' - '+str(row_drogi[2])+'km'
                            if row_drogi[1] == 'Vb':
                                a7 = row_drogi[1] +' - '+str(row_drogi[2])+'km'
                            drogi = [a1, a2, a3, a4, a5, a6, a7]
                            wypisz = ''
                            for droga in drogi:
                                if droga != '':
                                    wypisz = wypisz +droga+"; "
                    if wypisz == '':
                        add_tabrun(u'brak drogi wodnej', tab6.cell(0, 1).paragraphs[0])
                    else:
                        add_tabrun(wypisz, tab6.cell(0, 1).paragraphs[0])
                del cur_drogi

            #Wykorzystanie energii wody – elektrownie wodne
            if row[kol.index('MS_KOD')] not in lista_elek:
                add_tabrun(u'brak elektrowni wodnych', tab6.cell(1, 1).paragraphs[0])
            else:
                with arcpy.da.SearchCursor('elektrownie_opis', ['MS_KOD','Cnt_MS_KOD','MOC_MW']) as cur_elek:
                    for row_elek in cur_elek:
                        if  row[kol.index('MS_KOD')] == row_elek[0] and row_elek[2] == 0:
                            add_tabrun(u'Elektrownie wodne - ' + str(row_elek[1]) + '. Brak informacji o produkowanej mocy.', tab6.cell(1, 1).paragraphs[0])

                        elif row[kol.index('MS_KOD')] == row_elek[0] and row_elek[2] > 0:
                            add_tabrun(u'Elektrownie wodne - ' + unicode(row_elek[1]) + u'. Produkowana energia nie mniejsz niż ' + row_elek[2] + 'MW.', tab6.cell(1, 1).paragraphs[0])

                del cur_elek

            #Wielkość poborów  z wód powierzchniowych, dopuszczona w PWP, wg celu poboru [m3/rok]:
            for i in range(9):
                tab6.cell(2, 1).add_paragraph()

            b1 = b2 = b3 = b4 = b5= b6 = b7 = ''
            if row[kol.index('MS_KOD')] not in lista_pob:
                add_tabrun(u'brak poborów wód', tab6.cell(2, 1).paragraphs[0])
            else:
                with arcpy.da.SearchCursor('pobory', ['Kod_JCWP','_1','_2','_3','_4','_5','_6','_7']) as cur_pob:
                    for row_elek in cur_pob:
                        if  row[kol.index('MS_KOD')] == row_elek[0]:
                            b1 = "1) " + str(row_elek[1])
                            b2 = "2) " + str(row_elek[2])
                            b3 = "3) " + str(row_elek[3])
                            b4 = "4) " + str(row_elek[4])
                            b5 = "5) " + str(row_elek[5])
                            b6 = "6) " + str(row_elek[6])
                            b7 = "7) " + str(row_elek[7])
                del cur_pob

                add_tabrun(b1, tab6.cell(2, 1).paragraphs[2])
                add_tabrun(b2, tab6.cell(2, 1).paragraphs[3])
                add_tabrun(b3, tab6.cell(2, 1).paragraphs[4])
                add_tabrun(b4, tab6.cell(2, 1).paragraphs[5])
                add_tabrun(b5, tab6.cell(2, 1).paragraphs[6])
                add_tabrun(b6, tab6.cell(2, 1).paragraphs[7])
                add_tabrun(b7, tab6.cell(2, 1).paragraphs[8])

            #Obszar, na którym ryzyko powodziowe jest minimalizowane [ha]
            if row[kol.index('MS_KOD')] not in lista_och:
                add_tabrun(u'brak obszarów chronionych / brak zagrożenia powodzią', tab6.cell(3, 1).paragraphs[0])
            else:
                with arcpy.da.SearchCursor('Obsz_min_pow', ['MS_KOD','POW_HA']) as cur_och:
                    for row_och in cur_och:
                        if  row[kol.index('MS_KOD')] == row_och[0]:
                            add_tabrun(str(row_och[1]), tab6.cell(3, 1).paragraphs[0])
                del cur_och

            #Liczba ludzi zamieszkałych w chronionym obszarze
            if row[kol.index('MS_KOD')] not in lista_lch:
                add_tabrun(u'brak ludności chronionej / brak zagrożenia powodzią', tab6.cell(4, 1).paragraphs[0])
            else:
                with arcpy.da.SearchCursor('Ludnosc_chroniona', ['MS_KOD','Sum_L_OS_ZAM']) as cur_lch:
                    for row_lch in cur_lch:
                        if  row[kol.index('MS_KOD')] == row_lch[0] and row_lch[0] >0:
                            add_tabrun(str(row_lch[1]), tab6.cell(4, 1).paragraphs[0])
                        elif  row[kol.index('MS_KOD')] == row_lch[0] and row_lch[0] == 0:
                            add_tabrun(u'brak ludności zagrożonej powodzią', tab6.cell(4, 1).paragraphs[0])
                del cur_lch

            #Ważniejsze obiekty kulturowe, zagrażające środowisku w razie powodzi, które są chronione (obiekty kulturowe, zoo, cmentarze)
            if row[kol.index('MS_KOD')] not in lista_oc and row[kol.index('MS_KOD')] not in lista_zoo and row[kol.index('MS_KOD')] not in lista_cm:
                add_tabrun(u'brak obiektów', tab6.cell(5, 1).paragraphs[0])
            else:
                with arcpy.da.SearchCursor('ob_cenne', ['MS_KOD','Cnt_MS_KOD']) as cur_oc:
                    for row_oc in cur_oc:
                        if  row[kol.index('MS_KOD')] == row_oc[0]:
                            add_tabrun('obiekty cenne kulturowo - ' + str(row_oc[1]) + '; ', tab6.cell(5, 1).paragraphs[0])
                del cur_oc

                with arcpy.da.SearchCursor('zoo', ['MS_KOD','Cnt_MS_KOD']) as cur_zoo:
                    for row_zoo in cur_zoo:
                        if  row[kol.index('MS_KOD')] == row_zoo[0]:
                            add_tabrun('ogrody zoologiczne - ' + str(row_zoo[1]) + '; ', tab6.cell(5, 1).paragraphs[0])
                del cur_zoo

                with arcpy.da.SearchCursor('cmentarze', ['MS_KOD','Cnt_MS_KOD']) as cur_cm:
                    for row_cm in cur_cm:
                        if  row[kol.index('MS_KOD')] == row_cm[0]:
                            add_tabrun('cmentarze - ' + str(row_cm[1]) + '; ', tab6.cell(5, 1).paragraphs[0])
                del cur_cm

            #Ważniejsze obiekty infrastruktury społecznej, zagrażające środowisku w razie powodzi, które są chronione (oczyszczalnie, ujęcia wod, składowiska odpadow, kąpieliska)
            if row[kol.index('MS_KOD')] not in lista_kap and row[kol.index('MS_KOD')] not in lista_oczysz and row[kol.index('MS_KOD')] not in lista_uj and row[kol.index('MS_KOD')] not in lista_sklad:
                add_tabrun(u'brak obiektów', tab6.cell(6, 1).paragraphs[0])
            else:
                with arcpy.da.SearchCursor('ujecia_wod', ['MS_KOD','Cnt_MS_KOD']) as cur_uj:
                    for row_uj in cur_uj:
                        if  row[kol.index('MS_KOD')] == row_uj[0]:
                            add_tabrun(u'ujęcia wód - ' + str(row_uj[1]) + '; ', tab6.cell(6, 1).paragraphs[0])
                del cur_uj

                with arcpy.da.SearchCursor('oczyszczalnie', ['MS_KOD','Cnt_MS_KOD']) as cur_oczysz:
                    for row_oczysz in cur_oczysz:
                        if  row[kol.index('MS_KOD')] == row_oczysz[0]:
                            add_tabrun(u'oczyszczalnie - ' + str(row_oczysz[1]) + '; ', tab6.cell(6, 1).paragraphs[0])
                del cur_oczysz

                with arcpy.da.SearchCursor('skladowiska_odpadow', ['MS_KOD','Cnt_MS_KOD']) as cur_sklad:
                    for row_sklad in cur_sklad:
                        if  row[kol.index('MS_KOD')] == row_sklad[0]:
                            add_tabrun(u'składowiska odpadów - ' + str(row_sklad[1]) + '; ', tab6.cell(6, 1).paragraphs[0])
                del cur_sklad

                with arcpy.da.SearchCursor('kapieliska', ['MS_KOD','Cnt_MS_KOD']) as cur_kap:
                    for row_kap in cur_kap:
                        if  row[kol.index('MS_KOD')] == row_kap[0]:
                            add_tabrun(u'kąpieliska - ' + str(row_kap[1]) + '; ', tab6.cell(6, 1).paragraphs[0])
                del cur_kap

            #Ważniejsze obiekty gospodarcze, zagrażające środowisku w razie powodzi, które są chronione
            if row[kol.index('MS_KOD')] not in lista_zaklad:
                add_tabrun(u'brak obiektów', tab6.cell(7, 1).paragraphs[0])
            else:
                with arcpy.da.SearchCursor('zaklady_przem', ['MS_KOD','Cnt_MS_KOD']) as cur_zaklad:
                    for row_zaklad in cur_zaklad:
                        if  row[kol.index('MS_KOD')] == row_zaklad[0]:
                            add_tabrun(u'zakłady przemysłowe - ' + str(row_zaklad[1]) + '; ', tab6.cell(7, 1).paragraphs[0])
                del cur_zaklad

            #Obszary zmeliorowane [ha]
            if row[kol.index('MS_KOD')] not in lista_mel:
                add_tabrun(u'brak obszarów zmeliorowanych / brak obszarów zmeliorowanych oznaczonych w bazach referencyjnych', tab6.cell(8, 1).paragraphs[0])
            else:
                with arcpy.da.SearchCursor('ob_mel', ['MS_KOD','POW_HA']) as cur_mel:
                    for row_mel in cur_mel:
                        if  row[kol.index('MS_KOD')] == row_mel[0]:
                            add_tabrun(str(row_mel[1]), tab6.cell(8, 1).paragraphs[0])
                del cur_mel

            #Pobory kruszywa
            if row[kol.index('MS_KOD')] not in lista_krusz:
                add_tabrun(u'brak poborów kruszywa', tab6.cell(9, 1).paragraphs[0])
            else:
                with arcpy.da.SearchCursor('kruszywa_pob', ['MS_KOD','Sum_DL_OB']) as cur_krusz:
                    for row_krusz in cur_krusz:
                        if  row[kol.index('MS_KOD')] == row_krusz[0]:
                            add_tabrun(str(row_krusz[1]), tab6.cell(9, 1).paragraphs[0])
                del cur_krusz

            #Przerzuty wod
            if row[kol.index('MS_KOD')] not in lista_prze:
                add_tabrun(u'brak przerzutów', tab6.cell(10, 1).paragraphs[0])
            else:
                with arcpy.da.SearchCursor('przerzuty_wod', ['MS_KOD','Cel','Przerzut']) as cur_prze:
                    for row_prze in cur_prze:
                        if  row[kol.index('MS_KOD')] == row_prze[0]:
                            add_tabrun('cel przerzutu - ' + str(row_prze[1]) + ' - ' + str(row_prze[2]), tab6.cell(10, 1).paragraphs[0])
                del cur_prze

            #Wpływ górnictwa
            if row[kol.index('MS_KOD')] not in lista_gor:
                add_tabrun(u'brak wpływu górnictwa', tab6.cell(11, 1).paragraphs[0])
            else:
                with arcpy.da.SearchCursor('gornictwo', ['MS_KOD','OB_TG','POW_OG','ODW']) as cur_gor:
                    for row_gor in cur_gor:
                        if  row[kol.index('MS_KOD')] == row_gor[0]:
                            add_tabrun('1) ' + str(row_gor[1]), tab6.cell(11, 1).paragraphs[1])
                            add_tabrun('2) ' + str(row_gor[2]), tab6.cell(11, 1).paragraphs[2])
                            add_tabrun('3) ' + str(row_gor[3]), tab6.cell(11, 1).paragraphs[3])
                del cur_gor

            """Utworzenie karty informacyjnej"""
            document.save(os.path.join('zz_'+unicode(row[kol.index('ZZ')])+' '+str(row[kol.index('MS_KOD')])+'.docx'))
            print(str(row[kol.index('MS_KOD')])+'.docx')
            row[kol.index('KARTA')] = 999
            cur.updateRow(row)
del cur


def add_tabparagraph(tab_run, paragraph):

    run = paragraph.add_run(tab_run)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(10)
    return